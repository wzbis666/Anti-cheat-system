# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = 'docs/figures/usecase'

WHITELIST_USECASE_TABLE = [
    ['UC-6-01', '添加白名单', '管理员', '管理员已登录，玩家不在白名单中', '玩家被添加到白名单', '1.点击添加白名单按钮；2.输入玩家名称和UUID；3.填写添加原因；4.确认添加；5.系统创建白名单记录'],
    ['UC-6-02', '移除白名单', '管理员', '管理员已登录，玩家在白名单中', '玩家从白名单移除', '1.选择白名单玩家；2.点击移除按钮；3.确认移除操作；4.系统更新白名单状态'],
    ['UC-6-03', '查看白名单列表', '管理员', '管理员已登录', '显示所有白名单玩家', '1.进入白名单管理页面；2.系统查询白名单数据；3.展示白名单列表'],
    ['UC-6-04', '搜索白名单', '管理员', '管理员已登录', '显示匹配的白名单记录', '1.输入玩家名称；2.系统执行搜索；3.展示搜索结果'],
    ['UC-6-05', '查看白名单详情', '管理员', '管理员已登录', '显示白名单详细信息', '1.选择白名单记录；2.点击查看详情；3.系统展示白名单详情'],
    ['UC-6-06', '批量添加白名单', '管理员', '管理员已登录', '多个玩家被添加到白名单', '1.点击批量添加按钮；2.输入多个玩家信息；3.确认批量添加；4.系统批量创建记录'],
    ['UC-6-07', '批量移除白名单', '管理员', '管理员已登录', '多个玩家从白名单移除', '1.选择多个白名单玩家；2.点击批量移除按钮；3.确认批量移除；4.系统批量更新状态'],
    ['UC-6-08', '导出白名单', '管理员', '管理员已登录', '生成白名单数据文件', '1.点击导出按钮；2.系统生成数据文件；3.下载导出文件'],
]

def set_run_font(run, font_name='SimSun', font_size=12, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color

def set_cell_margins(cell, top=0, bottom=0, left=50, right=50):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_type, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_type}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_table(doc, headers, rows, title=None):
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        set_run_font(run, 'SimHei', 11, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(5)
    
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        row.height = Cm(0.8)
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_margins(hdr_cells[i], top=10, bottom=10, left=20, right=20)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, 'SimHei', 8, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
    
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            set_cell_margins(row_cells[col_idx], top=8, bottom=8, left=20, right=20)
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, 'SimSun', 8)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
    
    for col in table.columns:
        col.width = Cm(12.5)
    
    return table

def main():
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
    
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('白名单管理模块用例表')
    set_run_font(run, 'SimHei', 16, bold=True, color=RGBColor(0, 0, 0))
    title.paragraph_format.space_after = Pt(15)
    
    headers = ['用例编号', '用例名称', '参与者', '前置条件', '后置条件', '主要流程']
    
    add_table(doc, headers, WHITELIST_USECASE_TABLE, '表3-6 白名单管理模块用例表')
    
    filepath = os.path.join(OUTPUT_DIR, '表3-6_白名单管理模块用例表.docx')
    doc.save(filepath)
    print(f'已生成: {filepath}')
    print(f'完整路径: {os.path.abspath(filepath)}')

if __name__ == '__main__':
    main()
