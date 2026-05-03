# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_table(doc, title, headers, rows):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr_cells[i].paragraphs[0].runs[0].font.name = 'SimHei'
        hdr_cells[i].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(hdr_cells[i], 'E3F2FD')
    
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9)
            row_cells[col_idx].paragraphs[0].runs[0].font.name = 'SimSun'
            row_cells[col_idx].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()

def main():
    doc = Document()
    
    title = doc.add_heading('数据库表结构设计', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    headers = ['字段名', '数据类型', '约束', '说明']
    
    tables_data = [
        ('表1 管理员表 (admins)', [
            ('id', 'BIGINT', 'PK, AUTO_INCREMENT', '主键'),
            ('username', 'VARCHAR(50)', 'UNIQUE, NOT NULL', '用户名'),
            ('password', 'VARCHAR(255)', 'NOT NULL', '密码(BCrypt加密)'),
            ('nickname', 'VARCHAR(100)', '-', '昵称'),
            ('email', 'VARCHAR(100)', '-', '邮箱'),
            ('avatar', 'VARCHAR(255)', '-', '头像URL'),
            ('role', 'VARCHAR(20)', 'NOT NULL, DEFAULT ADMIN', '角色'),
            ('created_time', 'BIGINT', 'NOT NULL', '创建时间戳'),
            ('last_login_time', 'BIGINT', '-', '最后登录时间戳'),
            ('active', 'BOOLEAN', 'NOT NULL, DEFAULT TRUE', '是否激活'),
        ]),
        ('表2 用户表 (users)', [
            ('id', 'BIGINT', 'PK, AUTO_INCREMENT', '主键'),
            ('username', 'VARCHAR(50)', 'UNIQUE, NOT NULL', '用户名'),
            ('password', 'VARCHAR(255)', 'NOT NULL', '密码(BCrypt加密)'),
            ('nickname', 'VARCHAR(100)', '-', '昵称'),
            ('email', 'VARCHAR(100)', '-', '邮箱'),
            ('avatar', 'VARCHAR(255)', '-', '头像URL'),
            ('mc_username', 'VARCHAR(100)', '-', 'Minecraft用户名'),
            ('mc_uuid', 'VARCHAR(36)', '-', 'Minecraft UUID'),
            ('created_time', 'BIGINT', 'NOT NULL', '创建时间戳'),
            ('last_login_time', 'BIGINT', '-', '最后登录时间戳'),
            ('active', 'BOOLEAN', 'NOT NULL, DEFAULT TRUE', '是否激活'),
        ]),
        ('表3 玩家表 (players)', [
            ('id', 'BIGINT', 'PK, AUTO_INCREMENT', '主键'),
            ('player_name', 'VARCHAR(100)', 'UNIQUE, NOT NULL', '玩家名称'),
            ('uuid', 'VARCHAR(36)', 'UNIQUE, NOT NULL', 'Minecraft UUID'),
            ('risk_score', 'INT', 'NOT NULL, DEFAULT 0', '风险评分(0-100)'),
            ('last_seen', 'BIGINT', '-', '最后在线时间戳'),
            ('kick_count', 'INT', 'NOT NULL, DEFAULT 0', '被踢次数'),
        ]),
        ('表4 作弊记录表 (cheat_records)', [
            ('id', 'BIGINT', 'PK, AUTO_INCREMENT', '主键'),
            ('player_id', 'BIGINT', 'FK, NOT NULL', '关联玩家ID'),
            ('cheat_type', 'VARCHAR(50)', 'NOT NULL', '作弊类型'),
            ('detection_time', 'BIGINT', 'NOT NULL', '检测时间戳'),
            ('severity', 'INT', 'NOT NULL', '严重程度(1-10)'),
            ('details', 'TEXT', '-', '详细信息'),
        ]),
        ('表5 封禁记录表 (punishments)', [
            ('id', 'BIGINT', 'PK, AUTO_INCREMENT', '主键'),
            ('player_id', 'BIGINT', 'FK, NOT NULL', '关联玩家ID'),
            ('punishment_type', 'VARCHAR(20)', 'NOT NULL', '封禁类型'),
            ('punishment_time', 'BIGINT', 'NOT NULL', '封禁时间戳'),
            ('duration', 'BIGINT', '-', '封禁时长(毫秒)'),
            ('reason', 'VARCHAR(255)', '-', '封禁原因'),
            ('active', 'BOOLEAN', 'NOT NULL, DEFAULT TRUE', '是否生效'),
            ('unbanned_time', 'BIGINT', '-', '解封时间戳'),
            ('unbanned_by', 'VARCHAR(50)', '-', '解封操作人'),
        ]),
        ('表6 举报记录表 (reports)', [
            ('id', 'BIGINT', 'PK, AUTO_INCREMENT', '主键'),
            ('reporter_name', 'VARCHAR(100)', 'NOT NULL', '举报人名称'),
            ('reporter_uuid', 'VARCHAR(36)', 'NOT NULL', '举报人UUID'),
            ('reported_name', 'VARCHAR(100)', 'NOT NULL', '被举报人名称'),
            ('reported_uuid', 'VARCHAR(36)', 'NOT NULL', '被举报人UUID'),
            ('reason', 'VARCHAR(255)', 'NOT NULL', '举报原因'),
            ('report_type', 'VARCHAR(50)', '-', '举报类型'),
            ('report_time', 'BIGINT', 'NOT NULL', '举报时间戳'),
            ('status', 'VARCHAR(20)', 'NOT NULL, DEFAULT PENDING', '处理状态'),
            ('handled_by', 'VARCHAR(50)', '-', '处理人'),
            ('handled_time', 'BIGINT', '-', '处理时间戳'),
            ('result', 'VARCHAR(255)', '-', '处理结果'),
        ]),
        ('表7 白名单表 (whitelist)', [
            ('id', 'BIGINT', 'PK, AUTO_INCREMENT', '主键'),
            ('player_name', 'VARCHAR(100)', 'NOT NULL', '玩家名称'),
            ('uuid', 'VARCHAR(36)', 'UNIQUE, NOT NULL', 'Minecraft UUID'),
            ('reason', 'VARCHAR(255)', '-', '加入原因'),
            ('added_by', 'VARCHAR(50)', '-', '添加人'),
            ('added_time', 'BIGINT', 'NOT NULL', '添加时间戳'),
            ('active', 'BOOLEAN', 'NOT NULL, DEFAULT TRUE', '是否生效'),
        ]),
        ('表8 系统设置表 (system_settings)', [
            ('id', 'BIGINT', 'PK, AUTO_INCREMENT', '主键'),
            ('setting_key', 'VARCHAR(100)', 'UNIQUE, NOT NULL', '配置键'),
            ('setting_value', 'VARCHAR(500)', 'NOT NULL', '配置值'),
            ('description', 'VARCHAR(255)', '-', '配置说明'),
        ]),
    ]
    
    for title, rows in tables_data:
        create_table(doc, title, headers, rows)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('表9 作弊类型枚举值')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    enum_headers = ['枚举值', '中文名称', '说明']
    enum_rows = [
        ('KILLAURA', '杀戮光环', '自动攻击周围玩家'),
        ('FLY', '飞行作弊', '非法飞行或悬浮'),
        ('SPEED', '加速作弊', '移动速度异常'),
        ('XRAY', '透视作弊', '透视矿物或方块'),
        ('AUTOCLICKER', '自动点击', '自动点击鼠标'),
        ('REACH', '攻击距离作弊', '攻击距离超出正常范围'),
        ('OTHER', '其他', '其他作弊行为'),
    ]
    
    table = doc.add_table(rows=1 + len(enum_rows), cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(enum_headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr_cells[i].paragraphs[0].runs[0].font.name = 'SimHei'
        hdr_cells[i].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(hdr_cells[i], 'E3F2FD')
    
    for row_idx, row_data in enumerate(enum_rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9)
            row_cells[col_idx].paragraphs[0].runs[0].font.name = 'SimSun'
            row_cells[col_idx].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('表10 封禁类型枚举值')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    ban_enum_rows = [
        ('PERMANENT', '永久封禁', '无限期封禁玩家账号'),
        ('TEMPORARY', '临时封禁', '有期限的临时封禁'),
    ]
    
    table = doc.add_table(rows=1 + len(ban_enum_rows), cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(enum_headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr_cells[i].paragraphs[0].runs[0].font.name = 'SimHei'
        hdr_cells[i].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(hdr_cells[i], 'E3F2FD')
    
    for row_idx, row_data in enumerate(ban_enum_rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9)
            row_cells[col_idx].paragraphs[0].runs[0].font.name = 'SimSun'
            row_cells[col_idx].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('表11 举报状态枚举值')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    status_enum_rows = [
        ('PENDING', '待处理', '举报尚未处理'),
        ('RESOLVED', '已解决', '举报已确认并处理'),
        ('REJECTED', '已驳回', '举报被驳回'),
    ]
    
    table = doc.add_table(rows=1 + len(status_enum_rows), cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(enum_headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr_cells[i].paragraphs[0].runs[0].font.name = 'SimHei'
        hdr_cells[i].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(hdr_cells[i], 'E3F2FD')
    
    for row_idx, row_data in enumerate(status_enum_rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9)
            row_cells[col_idx].paragraphs[0].runs[0].font.name = 'SimSun'
            row_cells[col_idx].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save('docs/figures/数据库表结构设计.docx')
    print('Word文档已生成: docs/figures/数据库表结构设计.docx')

if __name__ == '__main__':
    main()
