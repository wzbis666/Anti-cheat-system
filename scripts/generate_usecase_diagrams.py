# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = 'docs/figures/usecase'

MODULES_DATA = {
    '用户管理模块': {
        'figure_num': '3-3',
        'table_num': '3-1',
        'actors': [
            {'name': '管理员', 'x': 60, 'y': 150},
            {'name': '普通用户', 'x': 60, 'y': 350},
        ],
        'usecases': [
            {'name': '用户登录', 'x': 250, 'y': 80, 'actors': ['管理员', '普通用户']},
            {'name': '用户注册', 'x': 250, 'y': 140, 'actors': ['普通用户']},
            {'name': '查看个人信息', 'x': 250, 'y': 200, 'actors': ['管理员', '普通用户']},
            {'name': '修改个人信息', 'x': 250, 'y': 260, 'actors': ['管理员', '普通用户']},
            {'name': '修改密码', 'x': 250, 'y': 320, 'actors': ['管理员', '普通用户']},
            {'name': '用户列表管理', 'x': 250, 'y': 380, 'actors': ['管理员']},
            {'name': '分配用户权限', 'x': 250, 'y': 440, 'actors': ['管理员']},
            {'name': '禁用/启用用户', 'x': 250, 'y': 500, 'actors': ['管理员']},
        ],
        'usecase_table': [
            ['UC-1-01', '用户登录', '管理员、普通用户', '用户未登录系统', '用户登录成功，获得JWT Token', '1.输入用户名和密码\n2.系统验证用户信息\n3.验证通过生成Token\n4.返回Token并跳转主页'],
            ['UC-1-02', '用户注册', '普通用户', '用户未注册账号', '创建新用户账号', '1.填写注册信息\n2.系统校验信息合法性\n3.密码加密存储\n4.创建用户记录'],
            ['UC-1-03', '查看个人信息', '管理员、普通用户', '用户已登录', '显示用户个人信息', '1.点击个人中心\n2.系统查询用户信息\n3.展示用户详细资料'],
            ['UC-1-04', '修改个人信息', '管理员、普通用户', '用户已登录', '更新用户个人信息', '1.进入个人信息编辑页\n2.修改需要更新的信息\n3.提交修改\n4.系统更新数据库'],
            ['UC-1-05', '修改密码', '管理员、普通用户', '用户已登录', '更新用户登录密码', '1.输入原密码和新密码\n2.系统验证原密码\n3.新密码加密存储\n4.更新密码成功'],
            ['UC-1-06', '用户列表管理', '管理员', '管理员已登录', '显示用户列表信息', '1.进入用户管理页面\n2.系统查询所有用户\n3.展示用户列表'],
            ['UC-1-07', '分配用户权限', '管理员', '管理员已登录', '更新用户权限', '1.选择目标用户\n2.修改用户角色\n3.确认修改\n4.更新用户权限'],
            ['UC-1-08', '禁用/启用用户', '管理员', '管理员已登录', '更新用户状态', '1.选择目标用户\n2.点击禁用/启用按钮\n3.确认操作\n4.更新用户状态'],
        ]
    },
    '玩家管理模块': {
        'figure_num': '3-4',
        'table_num': '3-2',
        'actors': [
            {'name': '管理员', 'x': 60, 'y': 280},
        ],
        'usecases': [
            {'name': '查看玩家列表', 'x': 250, 'y': 80, 'actors': ['管理员']},
            {'name': '搜索玩家', 'x': 250, 'y': 140, 'actors': ['管理员']},
            {'name': '查看玩家详情', 'x': 250, 'y': 200, 'actors': ['管理员']},
            {'name': '修改风险评分', 'x': 250, 'y': 260, 'actors': ['管理员']},
            {'name': '删除玩家', 'x': 250, 'y': 320, 'actors': ['管理员']},
            {'name': '按风险等级筛选', 'x': 250, 'y': 380, 'actors': ['管理员']},
            {'name': '查看玩家历史记录', 'x': 250, 'y': 440, 'actors': ['管理员']},
            {'name': '导出玩家数据', 'x': 250, 'y': 500, 'actors': ['管理员']},
        ],
        'usecase_table': [
            ['UC-2-01', '查看玩家列表', '管理员', '管理员已登录', '显示所有玩家列表', '1.进入玩家管理页面\n2.系统查询玩家数据\n3.分页展示玩家列表'],
            ['UC-2-02', '搜索玩家', '管理员', '管理员已登录', '显示匹配的玩家', '1.输入搜索关键词\n2.选择搜索条件\n3.系统执行搜索\n4.展示搜索结果'],
            ['UC-2-03', '查看玩家详情', '管理员', '管理员已登录', '显示玩家详细信息', '1.选择目标玩家\n2.点击查看详情\n3.系统查询玩家详细信息\n4.展示玩家详情页面'],
            ['UC-2-04', '修改风险评分', '管理员', '管理员已登录', '更新玩家风险评分', '1.选择目标玩家\n2.输入新的风险评分\n3.确认修改\n4.系统更新风险评分'],
            ['UC-2-05', '删除玩家', '管理员', '管理员已登录', '从系统中移除玩家', '1.选择目标玩家\n2.点击删除按钮\n3.确认删除操作\n4.系统删除玩家记录'],
            ['UC-2-06', '按风险等级筛选', '管理员', '管理员已登录', '显示指定风险等级的玩家', '1.选择风险等级\n2.系统筛选玩家\n3.展示筛选结果'],
            ['UC-2-07', '查看玩家历史记录', '管理员', '管理员已登录', '显示玩家历史作弊和封禁记录', '1.选择目标玩家\n2.点击查看历史\n3.系统查询历史记录\n4.展示历史记录列表'],
            ['UC-2-08', '导出玩家数据', '管理员', '管理员已登录', '生成玩家数据文件', '1.选择导出范围\n2.点击导出按钮\n3.系统生成数据文件\n4.下载导出文件'],
        ]
    },
    '作弊检测模块': {
        'figure_num': '3-5',
        'table_num': '3-3',
        'actors': [
            {'name': '管理员', 'x': 60, 'y': 200},
            {'name': 'MC玩家', 'x': 60, 'y': 400},
        ],
        'usecases': [
            {'name': '查看作弊记录列表', 'x': 250, 'y': 80, 'actors': ['管理员']},
            {'name': '搜索作弊记录', 'x': 250, 'y': 140, 'actors': ['管理员']},
            {'name': '查看作弊详情', 'x': 250, 'y': 200, 'actors': ['管理员']},
            {'name': '按作弊类型筛选', 'x': 250, 'y': 260, 'actors': ['管理员']},
            {'name': '按时间范围筛选', 'x': 250, 'y': 320, 'actors': ['管理员']},
            {'name': '实时告警接收', 'x': 250, 'y': 380, 'actors': ['管理员']},
            {'name': '作弊统计分析', 'x': 250, 'y': 440, 'actors': ['管理员']},
            {'name': '删除作弊记录', 'x': 250, 'y': 500, 'actors': ['管理员']},
        ],
        'usecase_table': [
            ['UC-3-01', '查看作弊记录列表', '管理员', '管理员已登录', '显示所有作弊记录', '1.进入作弊记录页面\n2.系统查询作弊数据\n3.分页展示作弊记录列表'],
            ['UC-3-02', '搜索作弊记录', '管理员', '管理员已登录', '显示匹配的作弊记录', '1.输入搜索条件\n2.系统执行搜索\n3.展示搜索结果'],
            ['UC-3-03', '查看作弊详情', '管理员', '管理员已登录', '显示作弊详细信息', '1.选择作弊记录\n2.点击查看详情\n3.系统展示作弊详情'],
            ['UC-3-04', '按作弊类型筛选', '管理员', '管理员已登录', '显示指定类型的作弊记录', '1.选择作弊类型\n2.系统筛选记录\n3.展示筛选结果'],
            ['UC-3-05', '按时间范围筛选', '管理员', '管理员已登录', '显示指定时间范围的作弊记录', '1.选择时间范围\n2.系统筛选记录\n3.展示筛选结果'],
            ['UC-3-06', '实时告警接收', '管理员', '管理员已登录，WebSocket已连接', '显示实时作弊告警', '1.系统接收作弊数据\n2.推送告警到前端\n3.前端显示告警通知'],
            ['UC-3-07', '作弊统计分析', '管理员', '管理员已登录', '显示作弊统计数据', '1.进入统计页面\n2.系统计算统计数据\n3.展示统计图表'],
            ['UC-3-08', '删除作弊记录', '管理员', '管理员已登录', '删除指定作弊记录', '1.选择作弊记录\n2.点击删除按钮\n3.确认删除\n4.系统删除记录'],
        ]
    },
    '封禁管理模块': {
        'figure_num': '3-6',
        'table_num': '3-4',
        'actors': [
            {'name': '管理员', 'x': 60, 'y': 200},
            {'name': 'MC玩家', 'x': 60, 'y': 420},
        ],
        'usecases': [
            {'name': '永久封禁玩家', 'x': 250, 'y': 80, 'actors': ['管理员']},
            {'name': '临时封禁玩家', 'x': 250, 'y': 140, 'actors': ['管理员']},
            {'name': '解封玩家', 'x': 250, 'y': 200, 'actors': ['管理员']},
            {'name': '查看封禁记录列表', 'x': 250, 'y': 260, 'actors': ['管理员']},
            {'name': '搜索封禁记录', 'x': 250, 'y': 320, 'actors': ['管理员']},
            {'name': '按封禁状态筛选', 'x': 250, 'y': 380, 'actors': ['管理员']},
            {'name': '查看封禁详情', 'x': 250, 'y': 440, 'actors': ['管理员']},
            {'name': '查询封禁状态', 'x': 250, 'y': 500, 'actors': ['MC玩家']},
        ],
        'usecase_table': [
            ['UC-4-01', '永久封禁玩家', '管理员', '管理员已登录，玩家未被永久封禁', '玩家被永久封禁', '1.选择目标玩家\n2.选择永久封禁类型\n3.填写封禁原因\n4.确认封禁\n5.系统创建封禁记录'],
            ['UC-4-02', '临时封禁玩家', '管理员', '管理员已登录，玩家未被封禁', '玩家被临时封禁', '1.选择目标玩家\n2.选择临时封禁类型\n3.设置封禁时长\n4.填写封禁原因\n5.确认封禁'],
            ['UC-4-03', '解封玩家', '管理员', '管理员已登录，玩家处于封禁状态', '玩家解除封禁', '1.选择已封禁玩家\n2.点击解封按钮\n3.确认解封操作\n4.系统更新封禁状态'],
            ['UC-4-04', '查看封禁记录列表', '管理员', '管理员已登录', '显示所有封禁记录', '1.进入封禁管理页面\n2.系统查询封禁数据\n3.分页展示封禁记录'],
            ['UC-4-05', '搜索封禁记录', '管理员', '管理员已登录', '显示匹配的封禁记录', '1.输入搜索条件\n2.系统执行搜索\n3.展示搜索结果'],
            ['UC-4-06', '按封禁状态筛选', '管理员', '管理员已登录', '显示指定状态的封禁记录', '1.选择封禁状态\n2.系统筛选记录\n3.展示筛选结果'],
            ['UC-4-07', '查看封禁详情', '管理员', '管理员已登录', '显示封禁详细信息', '1.选择封禁记录\n2.点击查看详情\n3.系统展示封禁详情'],
            ['UC-4-08', '查询封禁状态', 'MC玩家', '玩家在游戏中', '显示玩家封禁状态', '1.玩家发送查询请求\n2.系统查询封禁状态\n3.返回封禁状态信息'],
        ]
    },
    '举报管理模块': {
        'figure_num': '3-7',
        'table_num': '3-5',
        'actors': [
            {'name': '管理员', 'x': 60, 'y': 200},
            {'name': 'MC玩家', 'x': 60, 'y': 420},
        ],
        'usecases': [
            {'name': '提交举报', 'x': 250, 'y': 80, 'actors': ['MC玩家']},
            {'name': '查看举报列表', 'x': 250, 'y': 140, 'actors': ['管理员']},
            {'name': '搜索举报记录', 'x': 250, 'y': 200, 'actors': ['管理员']},
            {'name': '按处理状态筛选', 'x': 250, 'y': 260, 'actors': ['管理员']},
            {'name': '查看举报详情', 'x': 250, 'y': 320, 'actors': ['管理员']},
            {'name': '确认违规', 'x': 250, 'y': 380, 'actors': ['管理员']},
            {'name': '驳回举报', 'x': 250, 'y': 440, 'actors': ['管理员']},
            {'name': '查询举报状态', 'x': 250, 'y': 500, 'actors': ['MC玩家']},
        ],
        'usecase_table': [
            ['UC-5-01', '提交举报', 'MC玩家', '玩家在游戏中', '创建举报记录', '1.使用/report命令\n2.输入被举报玩家名称\n3.输入举报原因\n4.提交举报\n5.系统保存举报记录'],
            ['UC-5-02', '查看举报列表', '管理员', '管理员已登录', '显示所有举报记录', '1.进入举报管理页面\n2.系统查询举报数据\n3.分页展示举报列表'],
            ['UC-5-03', '搜索举报记录', '管理员', '管理员已登录', '显示匹配的举报记录', '1.输入搜索条件\n2.系统执行搜索\n3.展示搜索结果'],
            ['UC-5-04', '按处理状态筛选', '管理员', '管理员已登录', '显示指定状态的举报记录', '1.选择处理状态\n2.系统筛选记录\n3.展示筛选结果'],
            ['UC-5-05', '查看举报详情', '管理员', '管理员已登录', '显示举报详细信息', '1.选择举报记录\n2.点击查看详情\n3.系统展示举报详情'],
            ['UC-5-06', '确认违规', '管理员', '管理员已登录，举报待处理', '举报状态更新为已确认', '1.选择待处理举报\n2.查看举报详情\n3.点击确认违规\n4.选择是否封禁玩家\n5.记录处理结果'],
            ['UC-5-07', '驳回举报', '管理员', '管理员已登录，举报待处理', '举报状态更新为已驳回', '1.选择待处理举报\n2.查看举报详情\n3.点击驳回举报\n4.填写驳回原因\n5.记录处理结果'],
            ['UC-5-08', '查询举报状态', 'MC玩家', '玩家已提交举报', '显示举报处理状态', '1.发送查询请求\n2.系统查询举报状态\n3.返回处理状态信息'],
        ]
    },
    '白名单管理模块': {
        'figure_num': '3-8',
        'table_num': '3-6',
        'actors': [
            {'name': '管理员', 'x': 60, 'y': 280},
        ],
        'usecases': [
            {'name': '添加白名单', 'x': 250, 'y': 80, 'actors': ['管理员']},
            {'name': '移除白名单', 'x': 250, 'y': 140, 'actors': ['管理员']},
            {'name': '查看白名单列表', 'x': 250, 'y': 200, 'actors': ['管理员']},
            {'name': '搜索白名单', 'x': 250, 'y': 260, 'actors': ['管理员']},
            {'name': '查看白名单详情', 'x': 250, 'y': 320, 'actors': ['管理员']},
            {'name': '批量添加白名单', 'x': 250, 'y': 380, 'actors': ['管理员']},
            {'name': '批量移除白名单', 'x': 250, 'y': 440, 'actors': ['管理员']},
            {'name': '导出白名单', 'x': 250, 'y': 500, 'actors': ['管理员']},
        ],
        'usecase_table': [
            ['UC-6-01', '添加白名单', '管理员', '管理员已登录，玩家不在白名单中', '玩家被添加到白名单', '1.点击添加白名单按钮\n2.输入玩家名称和UUID\n3.填写添加原因\n4.确认添加\n5.系统创建白名单记录'],
            ['UC-6-02', '移除白名单', '管理员', '管理员已登录，玩家在白名单中', '玩家从白名单移除', '1.选择白名单玩家\n2.点击移除按钮\n3.确认移除操作\n4.系统更新白名单状态'],
            ['UC-6-03', '查看白名单列表', '管理员', '管理员已登录', '显示所有白名单玩家', '1.进入白名单管理页面\n2.系统查询白名单数据\n3.展示白名单列表'],
            ['UC-6-04', '搜索白名单', '管理员', '管理员已登录', '显示匹配的白名单记录', '1.输入玩家名称\n2.系统执行搜索\n3.展示搜索结果'],
            ['UC-6-05', '查看白名单详情', '管理员', '管理员已登录', '显示白名单详细信息', '1.选择白名单记录\n2.点击查看详情\n3.系统展示白名单详情'],
            ['UC-6-06', '批量添加白名单', '管理员', '管理员已登录', '多个玩家被添加到白名单', '1.点击批量添加按钮\n2.输入多个玩家信息\n3.确认批量添加\n4.系统批量创建记录'],
            ['UC-6-07', '批量移除白名单', '管理员', '管理员已登录', '多个玩家从白名单移除', '1.选择多个白名单玩家\n2.点击批量移除按钮\n3.确认批量移除\n4.系统批量更新状态'],
            ['UC-6-08', '导出白名单', '管理员', '管理员已登录', '生成白名单数据文件', '1.点击导出按钮\n2.系统生成数据文件\n3.下载导出文件'],
        ]
    },
    '数据统计模块': {
        'figure_num': '3-9',
        'table_num': '3-7',
        'actors': [
            {'name': '管理员', 'x': 60, 'y': 280},
        ],
        'usecases': [
            {'name': '查看概览统计', 'x': 250, 'y': 80, 'actors': ['管理员']},
            {'name': '查看作弊类型分布', 'x': 250, 'y': 140, 'actors': ['管理员']},
            {'name': '查看检测趋势', 'x': 250, 'y': 200, 'actors': ['管理员']},
            {'name': '查看风险等级分布', 'x': 250, 'y': 260, 'actors': ['管理员']},
            {'name': '查看实时告警', 'x': 250, 'y': 320, 'actors': ['管理员']},
            {'name': '刷新统计数据', 'x': 250, 'y': 380, 'actors': ['管理员']},
            {'name': '导出统计报表', 'x': 250, 'y': 440, 'actors': ['管理员']},
            {'name': '自定义时间范围统计', 'x': 250, 'y': 500, 'actors': ['管理员']},
        ],
        'usecase_table': [
            ['UC-7-01', '查看概览统计', '管理员', '管理员已登录', '显示系统概览数据', '1.进入仪表盘页面\n2.系统计算统计数据\n3.展示概览统计卡片'],
            ['UC-7-02', '查看作弊类型分布', '管理员', '管理员已登录', '显示作弊类型分布图', '1.进入统计页面\n2.系统统计各类型数量\n3.展示饼图'],
            ['UC-7-03', '查看检测趋势', '管理员', '管理员已登录', '显示检测趋势图', '1.进入统计页面\n2.系统统计时间趋势\n3.展示折线图'],
            ['UC-7-04', '查看风险等级分布', '管理员', '管理员已登录', '显示风险等级分布图', '1.进入统计页面\n2.系统统计风险分布\n3.展示分布图'],
            ['UC-7-05', '查看实时告警', '管理员', '管理员已登录', '显示实时告警列表', '1.进入仪表盘页面\n2.系统推送实时告警\n3.展示告警列表'],
            ['UC-7-06', '刷新统计数据', '管理员', '管理员已登录', '更新统计数据', '1.点击刷新按钮\n2.系统重新计算统计\n3.更新展示数据'],
            ['UC-7-07', '导出统计报表', '管理员', '管理员已登录', '生成统计报表文件', '1.选择导出范围\n2.点击导出按钮\n3.系统生成报表\n4.下载报表文件'],
            ['UC-7-08', '自定义时间范围统计', '管理员', '管理员已登录', '显示指定时间范围的统计数据', '1.选择时间范围\n2.系统计算统计数据\n3.展示统计结果'],
        ]
    },
    '系统设置模块': {
        'figure_num': '3-10',
        'table_num': '3-8',
        'actors': [
            {'name': '管理员', 'x': 60, 'y': 280},
        ],
        'usecases': [
            {'name': '查看系统设置', 'x': 250, 'y': 80, 'actors': ['管理员']},
            {'name': '修改系统参数', 'x': 250, 'y': 140, 'actors': ['管理员']},
            {'name': '设置自动封禁阈值', 'x': 250, 'y': 200, 'actors': ['管理员']},
            {'name': '配置检测规则', 'x': 250, 'y': 260, 'actors': ['管理员']},
            {'name': '查看系统信息', 'x': 250, 'y': 320, 'actors': ['管理员']},
            {'name': '查看操作日志', 'x': 250, 'y': 380, 'actors': ['管理员']},
            {'name': '系统数据初始化', 'x': 250, 'y': 440, 'actors': ['管理员']},
            {'name': '备份系统数据', 'x': 250, 'y': 500, 'actors': ['管理员']},
        ],
        'usecase_table': [
            ['UC-8-01', '查看系统设置', '管理员', '管理员已登录', '显示系统设置列表', '1.进入系统设置页面\n2.系统查询设置数据\n3.展示设置列表'],
            ['UC-8-02', '修改系统参数', '管理员', '管理员已登录', '更新系统参数', '1.选择需要修改的参数\n2.输入新的参数值\n3.确认修改\n4.系统更新参数'],
            ['UC-8-03', '设置自动封禁阈值', '管理员', '管理员已登录', '更新自动封禁阈值', '1.进入阈值设置页面\n2.输入风险评分阈值\n3.确认设置\n4.系统更新阈值'],
            ['UC-8-04', '配置检测规则', '管理员', '管理员已登录', '更新检测规则配置', '1.进入检测规则配置页面\n2.修改规则参数\n3.确认配置\n4.系统更新规则'],
            ['UC-8-05', '查看系统信息', '管理员', '管理员已登录', '显示系统版本和状态信息', '1.进入系统信息页面\n2.系统收集系统信息\n3.展示系统信息'],
            ['UC-8-06', '查看操作日志', '管理员', '管理员已登录', '显示操作日志列表', '1.进入日志查看页面\n2.系统查询日志数据\n3.展示日志列表'],
            ['UC-8-07', '系统数据初始化', '管理员', '管理员已登录', '系统数据重置', '1.进入初始化页面\n2.确认初始化操作\n3.系统执行初始化\n4.返回初始化结果'],
            ['UC-8-08', '备份系统数据', '管理员', '管理员已登录', '生成系统数据备份', '1.点击备份按钮\n2.系统执行数据备份\n3.生成备份文件\n4.下载备份文件'],
        ]
    },
}

def create_usecase_diagram(module_name, data):
    actors_xml = ""
    usecases_xml = ""
    connections_xml = ""
    actor_id = 100
    usecase_id = 200
    conn_id = 300
    
    actor_positions = {}
    
    for actor in data['actors']:
        actor_positions[actor['name']] = actor_id
        actors_xml += f'''        <mxCell id="{actor_id}" value="{actor['name']}" style="shape=umlActor;verticalLabelPosition=bottom;verticalAlign=top;html=1;outlineConnect=0;strokeColor=#000000;fontColor=#000000;fontSize=12;" vertex="1" parent="1">
          <mxGeometry x="{actor['x']}" y="{actor['y']}" width="40" height="80" as="geometry"/>
        </mxCell>
'''
        actor_id += 1
    
    for usecase in data['usecases']:
        usecase_id_str = str(usecase_id)
        usecases_xml += f'''        <mxCell id="{usecase_id}" value="{usecase['name']}" style="ellipse;whiteSpace=wrap;html=1;fillColor=#ffffff;strokeColor=#000000;fontColor=#000000;fontSize=11;" vertex="1" parent="1">
          <mxGeometry x="{usecase['x']}" y="{usecase['y']}" width="120" height="45" as="geometry"/>
        </mxCell>
'''
        for actor_name in usecase['actors']:
            actor_id_ref = actor_positions.get(actor_name, 100)
            connections_xml += f'''        <mxCell id="{conn_id}" style="endArrow=open;html=1;endSize=8;strokeColor=#000000;" edge="1" parent="1" source="{actor_id_ref}" target="{usecase_id_str}">
          <mxGeometry relative="1" as="geometry"/>
        </mxCell>
'''
            conn_id += 1
        
        usecase_id += 1
    
    system_boundary_height = max([uc['y'] for uc in data['usecases']]) + 100
    
    return f'''<mxfile host="app.diagrams.net" agent="Mozilla/5.0" version="21.0.0">
  <diagram id="diagram1" name="Page-1">
    <mxGraphModel dx="900" dy="700" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="827" pageHeight="1169" math="0" shadow="0" background="#ffffff">
      <root>
        <mxCell id="0"/>
        <mxCell id="1" parent="0"/>
        <mxCell id="2" value="{module_name}" style="rounded=0;whiteSpace=wrap;html=1;fillColor=#ffffff;strokeColor=#000000;fontSize=14;fontStyle=1;verticalAlign=top;fontColor=#000000;align=center;" vertex="1" parent="1">
          <mxGeometry x="200" y="30" width="220" height="{system_boundary_height}" as="geometry"/>
        </mxCell>
{actors_xml}{usecases_xml}{connections_xml}      </root>
    </mxGraphModel>
  </diagram>
</mxfile>'''

def set_run_font(run, font_name='SimSun', font_size=12, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color

def set_cell_margins(cell, top=0, bottom=0, left=50, right=50):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_type, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_type}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_table(doc, headers, rows, title=None):
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        set_run_font(run, 'SimHei', 11, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(5)
    
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        row.height = Cm(0.8)
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_margins(hdr_cells[i], top=10, bottom=10, left=20, right=20)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, 'SimHei', 8, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
    
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            text = str(cell_text).replace('\n', '；')
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx], top=8, bottom=8, left=20, right=20)
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, 'SimSun', 8)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
    
    for col in table.columns:
        col.width = Cm(12.5)
    
    return table

def create_usecase_tables_doc():
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('系统功能模块用例表汇总')
    set_run_font(run, 'SimHei', 18, bold=True, color=RGBColor(0, 0, 0))
    title.paragraph_format.space_after = Pt(20)
    
    headers = ['用例编号', '用例名称', '参与者', '前置条件', '后置条件', '主要流程']
    
    for module_name, data in MODULES_DATA.items():
        module_title = doc.add_paragraph()
        run = module_title.add_run(f'{module_name}用例表（表{data["table_num"]}）')
        set_run_font(run, 'SimHei', 14, bold=True)
        module_title.paragraph_format.space_before = Pt(15)
        module_title.paragraph_format.space_after = Pt(10)
        
        add_table(doc, headers, data['usecase_table'])
        
        doc.add_paragraph()
    
    return doc

def main():
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
    
    print('=' * 60)
    print('生成用例图和用例表')
    print('=' * 60)
    
    print('\n[1] 生成用例图（draw.io XML格式）')
    print('-' * 40)
    for module_name, data in MODULES_DATA.items():
        filename = f'图{data["figure_num"]}_{module_name}用例图.drawio'
        filepath = os.path.join(OUTPUT_DIR, filename)
        xml_content = create_usecase_diagram(module_name, data)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(xml_content)
        print(f'  已生成: {filepath}')
    
    print('\n[2] 生成用例表（Word文档）')
    print('-' * 40)
    doc = create_usecase_tables_doc()
    doc_path = os.path.join(OUTPUT_DIR, '系统功能模块用例表汇总.docx')
    doc.save(doc_path)
    print(f'  已生成: {doc_path}')
    
    print('\n[3] 生成单独的用例表（Word文档）')
    print('-' * 40)
    headers = ['用例编号', '用例名称', '参与者', '前置条件', '后置条件', '主要流程']
    for module_name, data in MODULES_DATA.items():
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
        
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f'{module_name}用例表')
        set_run_font(run, 'SimHei', 16, bold=True, color=RGBColor(0, 0, 0))
        title.paragraph_format.space_after = Pt(15)
        
        add_table(doc, headers, data['usecase_table'], f'表{data["table_num"]} {module_name}用例表')
        
        filename = f'表{data["table_num"]}_{module_name}用例表.docx'
        filepath = os.path.join(OUTPUT_DIR, filename)
        doc.save(filepath)
        print(f'  已生成: {filepath}')
    
    print('\n' + '=' * 60)
    print('生成完成！')
    print(f'用例图: {len(MODULES_DATA)} 个 draw.io 文件')
    print(f'用例表: 1 个汇总文档 + {len(MODULES_DATA)} 个单独文档')
    print(f'输出目录: {os.path.abspath(OUTPUT_DIR)}')
    print('=' * 60)

if __name__ == '__main__':
    main()
